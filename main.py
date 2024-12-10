import os
import sys

import docx
from playsound import playsound

import shutil
from docx import Document
# from docx2pdf import convert
from spire.doc import *
from spire.doc.common import *

import errorWindow
import helpWindow
import choiceWindow
import sounds
import tables

import sqlite3

from PyQt6 import uic
from PyQt6.QtWidgets import QApplication, QDialog


# создание основного рабочего окна программы
class MainWidget(QDialog):
    def __init__(self):
        # инициализация значений
        super().__init__()
        uic.loadUi('mainWindow.ui', self)

        # привязка функций к методам
        self.createrFileButton.clicked.connect(self.create_file)
        self.helpButton.clicked.connect(self.open_help_window)
        self.choiceDocument.clicked.connect(self.choice_file)

        # установление значений для окон с выбором вариантов
        self.comboFormatBox.addItems(tables.FILE_FORMATS)
        self.comboTasksBox.addItems(tables.AREAS_OF_ACTIVITY)

        # индекс выбранного документа
        # !!!!!!!!!!!!!!!!!!!!!!!!!!!
        self.index_file = '00-000'
        self.documentName.setPlainText(tables.ALL_DOCUMENTS[self.index_file])

    # создание сайта
    def create_file(self):
        # получаем введённые пользователи значения из интерфейса
        no_choice = 'не выбранно'
        task = self.comboTasksBox.currentText()
        format = self.comboFormatBox.currentText()
        file = self.documentName.toPlainText()

        # проверяем, что все данные введены
        if task == no_choice:
            # playsound(sounds.common_path + sounds.sound_error_window)
            self.open_error_window('Пожалуйста, выберите\nсферу')
        elif format == no_choice:
            # playsound(sounds.common_path + sounds.sound_error_window)
            self.open_error_window('Пожалуйста, выберите\nформат файла')
        # elif file == 'Документ не выбран':
        elif self.index_file == '00-000':
            self.open_error_window('Пожалуйста, выберите\nдокумент')

        # создаём файл
        else:
            folder = self.create_folder(task)
            file_name = f'{file}.{self.comboFormatBox.currentText()}'

            new_text = self.get_text(self.index_file)
            new_index = self.get_key(tables.ALL_DOCUMENTS, new_text)

            # код для конверции файла в разные форматы
            if self.comboFormatBox.currentText() == 'pdf':
                docx_file = Document()
                docx_file.LoadFromFile(f'Documents\\{self.index_file}.docx')

                parameter = ToPdfParameterList()
                parameter.DisableLink = True
                parameter.IsEmbeddedAllFonts = True

                docx_file.SaveToFile(f'{folder}\\{tables.ALL_DOCUMENTS[self.index_file]}.pdf', parameter)
                docx_file.Close()

            elif self.comboFormatBox.currentText() == 'docx':
                docx_file = f'Documents\\{self.index_file}.docx'
                copy_file = folder + '\\' + f'{self.index_file}.docx'
                new_file = folder + '\\' + f'{tables.ALL_DOCUMENTS[self.index_file]}.docx'

                shutil.copy(docx_file, copy_file)
                os.rename(copy_file, new_file)

            elif self.comboFormatBox.currentText() == 'txt':
                docx_file = docx.Document(f'Documents\\{self.index_file}.docx')
                new_file = open(f'{folder}\\{file_name}', 'w')

                for paragraph in docx_file.paragraphs:
                    new_file.write(paragraph.text)
                new_file.close()

            # сбрасываем значения
            self.index_file = '00-000'
            print(new_index)
            self.documentName.setPlainText(tables.ALL_DOCUMENTS[self.index_file])

    # получаем значение из базы данных
    def get_text(self, code):
        db = sqlite3.connect('data.db')
        cur = db.cursor()
        cur.execute(f'SELECT text FROM documents WHERE code = "{code}"')
        text = cur.fetchone()
        return text

    # получаем ключ словаря через значен ие ключа
    def get_key(self, d, value):
        for k, v in d.items():
            if v == value:
                return k

    # создание папки
    def create_folder(self, task):
        main_folder = r'C:\Users\Admin\Desktop\DocumentsPy'
        task_folder = main_folder + '\\' + task

        if not os.path.exists(main_folder):
            os.makedirs(main_folder)
        if not os.path.exists(task_folder):
            os.makedirs(task_folder)

        return task_folder

    # открытие окна для выбора документа
    def choice_file(self):
        text_buttons = None
        index_buttons = None

        # проверка входных значений пользователя
        if self.comboTasksBox.currentText() == 'не выбранно':
            self.open_error_window('Пожалуйста, выберите\nсферу')
            return
        elif self.comboFormatBox.currentText() == 'не выбранно':
            self.open_error_window('Пожалуйста, выберите\nформат')
            return

        # установка характеристик для объектов
        if self.comboTasksBox.currentText() == 'Юрист':
            text_buttons = ['Договор на \nоказание \nюридических \nуслуг',
                            'Доверенность \nна представление \nинтересов \nфизического лица в \nсудебных органах']
            index_buttons = ['01-001', '01-002']
        elif self.comboTasksBox.currentText() == 'Педагог':
            text_buttons = ['Договор об \nоказании платных \nобразовательных \nуслуг',
                            'Договор оказания \nобразовательных \nуслуг по программе \nповышения \nквалификации']
            index_buttons = ['02-001', '02-002']
        elif self.comboTasksBox.currentText() == 'Бугалтер':
            text_buttons = ['Договор на \nоказание \nбугалтерских \nуслуг',
                            'Договор на \nоказание услуг по \nведению \nбугалтерского \nсчёта']
            index_buttons = ['03-001', '03-002']
        elif self.comboTasksBox.currentText() == 'Другое':
            text_buttons = ['Договор \nкупли-продажи \nнедвижимости',
                            'Договор \nкупли-продажи \nтранспортного \nсредства']
            index_buttons = ['04-001', '04-002']
        elif self.comboTasksBox.currentText() == 'не выбранно':
            self.open_error_window('Пожалуйста, выберите\nсферу')
            return

        # создание и открытие окна
        choice_window = choiceWindow.ChoiceWidget(text_buttons, index_buttons)
        choice_window.exec()
        self.update_document(choice_window.index)

    # обновляем значение выбранного файла
    def update_document(self, new_index):
        self.index_file = new_index
        self.documentName.setPlainText(tables.ALL_DOCUMENTS[self.index_file])

    # вызов окна с ошибкой
    def open_error_window(self, name_error):
        error_window = errorWindow.ErrorWidget(name_error)
        error_window.exec()

    # открытие окна помощи
    def open_help_window(self):
        help_window = helpWindow.HelpWidget()
        help_window.exec()


# запуск программы
if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = MainWidget()
    ex.show()
    sys.exit(app.exec())
